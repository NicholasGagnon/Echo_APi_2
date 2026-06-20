import io
import os
import re
from datetime import datetime
from flask import jsonify, send_file, request

# --- LIBRAIRIES D'EXPORTATION ---
# Ces bibliothèques doivent être installées sur ton serveur Flask via ton requirements.txt
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    epub = None

def generate_pdf_buffer(title: str, html_content: str) -> io.BytesIO:
    """
    Génère un vrai fichier PDF vectoriel à partir d'un template HTML et CSS.
    """
    if pisa is None:
        raise ImportError("La bibliothèque 'xhtml2pdf' n'est pas installée sur le serveur Flask.")

    styled_html = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {{
                size: a4;
                margin: 2cm;
                @bottom-right {{
                    content: "Page " counter(page);
                    font-family: sans-serif;
                    font-size: 9pt;
                    color: #71717a;
                }}
            }}
            body {{
                font-family: sans-serif;
                color: #18181b;
                line-height: 1.6;
                font-size: 11pt;
            }}
            h1 {{
                font-family: sans-serif;
                color: #0f172a;
                font-size: 24pt;
                border-bottom: 1px solid #e4e4e7;
                padding-bottom: 8px;
                margin-bottom: 20px;
            }}
            p {{
                margin-bottom: 15px;
                text-align: justify;
            }}
            .metadata {{
                font-size: 9pt;
                color: #71717a;
                margin-bottom: 30px;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        <div class="metadata">Généré par Echo AI — {datetime.now().strftime('%d/%m/%Y à %H:%M')}</div>
        <div class="content">
            {html_content}
        </div>
    </body>
    </html>
    """
    
    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(styled_html), dest=pdf_buffer)
    
    if pisa_status.err:
        raise Exception("Erreur lors de la compilation du PDF par l'assistant pisa.")
        
    pdf_buffer.seek(0)
    return pdf_buffer

def generate_docx_buffer(title: str, text_content: str) -> io.BytesIO:
    """
    Génère un VRAI fichier Microsoft Word .docx structuré au format OpenXML.
    """
    if Document is None:
        raise ImportError("La bibliothèque 'python-docx' n'est pas installée sur le serveur Flask.")

    doc = Document()
    
    # Configuration des marges de la page
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Ajout du Titre Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    # Métadonnées
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(f"Document généré par Echo AI — {datetime.now().strftime('%d/%m/%Y à %H:%M')}\n")
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(9.5)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(113, 113, 122)
    
    doc.add_paragraph("__________________________________________________________________")

    # Injection du texte standard par paragraphe
    paragraphs = text_content.split('\n')
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(p_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(24, 24, 27)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def generate_epub_buffer(title: str, text_content: str) -> io.BytesIO:
    """
    Génère un livre numérique EPUB v3 pour les liseuses.
    """
    if epub is None:
        raise ImportError("La bibliothèque 'EbookLib' n'est pas installée sur le serveur Flask.")

    book = epub.EpubBook()

    book.set_identifier(f"echo-ai-document-{int(datetime.now().timestamp())}")
    book.set_title(title)
    book.set_language('fr')
    book.add_author('Echo AI')

    c1 = epub.EpubHtml(title=title, file_name='chapitre_1.xhtml', lang='fr')
    
    paragraphs_xhtml = ""
    for p in text_content.split('\n'):
        if p.strip():
            paragraphs_xhtml += f"<p>{p.strip()}</p>\n"

    c1.content = f"""
    <?xml version="1.0" encoding="utf-8"?>
    <!DOCTYPE html>
    <html xmlns="http://www.w3.org/1999/xhtml" xml:lang="fr">
    <head>
        <title>{title}</title>
        <style>
            body {{ font-family: Georgia, serif; line-height: 1.5; padding: 5%; }}
            h1 {{ color: #0f172a; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
            p {{ text-align: justify; margin-bottom: 1em; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        {paragraphs_xhtml}
    </body>
    </html>
    """

    book.add_item(c1)
    book.toc = (epub.Link('chapitre_1.xhtml', title, 'chapitre1'),)
    book.spine = ['nav', c1]
    book.add_item(epub.EpubNav())
    book.add_item(epub.EpubNcx())

    epub_buffer = io.BytesIO()
    epub.write_epub(epub_buffer, book, {})
    epub_buffer.seek(0)
    return epub_buffer

def handle_file_download(format_type: str, title: str, content: str):
    """
    Fonction principale pour traiter la demande et renvoyer le fichier.
    """
    try:
        format_type = format_type.lower().strip()
        safe_filename = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        safe_filename = safe_filename.replace(' ', '_') or "document_export"

        if format_type == "pdf":
            html_formatted = content.replace('\n', '<br>')
            buffer = generate_pdf_buffer(title, html_formatted)
            return send_file(
                buffer,
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"{safe_filename}.pdf"
            )

        elif format_type == "docx":
            buffer = generate_docx_buffer(title, content)
            return send_file(
                buffer,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"{safe_filename}.docx"
            )

        elif format_type == "epub":
            buffer = generate_epub_buffer(title, content)
            return send_file(
                buffer,
                mimetype="application/epub+zip",
                as_attachment=True,
                download_name=f"{safe_filename}.epub"
            )

        elif format_type == "txt":
            clean_text = re.sub(r'<[^>]+>', '', content)
            clean_text = clean_text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
            
            buffer = io.BytesIO(clean_text.encode('utf-8'))
            return send_file(
                buffer,
                mimetype="text/plain",
                as_attachment=True,
                download_name=f"{safe_filename}.txt"
            )

        else:
            return jsonify({"error": "Format d'exportation non supporté."}), 400

    except Exception as e:
        print(f"❌ [EXPORT ERROR] Échec de la génération {format_type} : {e}")
        return jsonify({"error": f"Impossible de générer le fichier : {str(e)}"}), 500