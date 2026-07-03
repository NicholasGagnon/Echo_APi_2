import os
import io
import re
import json
import base64
import threading
import requests as _requests_lib
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google import genai
from google.genai import types
from openai import OpenAI
from dotenv import load_dotenv

from prompts import generate_system_prompt

load_dotenv()

app = Flask(__name__)
CORS(app)

from site2 import site2_bp
app.register_blueprint(site2_bp)

# ── MODÈLES ────────────────────────────────────────────────────────────────────
MODELS = {
    "gemini_paid_founder":  "gemini-3.5-flash",
    "gemini_paid_ultra":    "gemini-3.1-flash-lite",
    "gemini_paid_standard": "gemini-2.5-flash-lite",
    "mistral":    "amazon/nova-lite-v1",
    "hy3":        "tencent/hy3-preview",
    "owl":        "openrouter/owl-alpha",
    "llama":      "meta-llama/llama-3.3-70b-instruct",
    "ling":       "inclusionai/ling-2.6-flash",
    "deepseek":   "deepseek-v4-flash",
    "glm":        "GLM-4.5-Air",
}

FREE_MAX_TOKENS = 2_500
PAID_MAX_TOKENS = 5_000

# ── CLÉS & CLIENTS ─────────────────────────────────────────────────────────────
API_KEY_PAID          = os.getenv("API_KEY_PAID", "").strip()
OPENROUTER_API_KEY    = os.getenv("OPENROUTER_API_KEY", "").strip()
DEEPSEEK_API_KEY      = os.getenv("DEEPSEEK_API_KEY", "").strip()
ZAI_API_KEY           = os.getenv("ZAI_API_KEY", "").strip()
CLOUDFLARE_ACCOUNT_ID = os.getenv("CLOUDFLARE_ACCOUNT_ID", "").strip()

client_gemini_paid = genai.Client(api_key=API_KEY_PAID) if API_KEY_PAID else None

import httpx
_shared_http_client = httpx.Client(timeout=35.0)

client_openrouter = (
    OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY, http_client=_shared_http_client)
    if OPENROUTER_API_KEY else None
)
client_deepseek = (
    OpenAI(base_url="https://api.deepseek.com", api_key=DEEPSEEK_API_KEY, http_client=_shared_http_client)
    if DEEPSEEK_API_KEY else None
)
client_zai = (
    OpenAI(base_url="https://api.z.ai/api/paas/v4", api_key=ZAI_API_KEY, http_client=_shared_http_client)
    if ZAI_API_KEY else None
)

_cf_session = _requests_lib.Session()

# ── FAILOVER GLOBAL ────────────────────────────────────────────────────────────
GLOBAL_FAILOVER_MEMORY_COUNT = 0
MAX_FREE_FAILOVERS = 200

def get_failover_count():
    global GLOBAL_FAILOVER_MEMORY_COUNT
    return GLOBAL_FAILOVER_MEMORY_COUNT

def increment_failover_count():
    global GLOBAL_FAILOVER_MEMORY_COUNT
    GLOBAL_FAILOVER_MEMORY_COUNT += 1
    print(f"[FAILOVER] {GLOBAL_FAILOVER_MEMORY_COUNT} / {MAX_FREE_FAILOVERS}")

# ── LOCK MODÈLES — 30s ─────────────────────────────────────────────────────────
MODELS_LOCK_REGISTRY = {}
_lock_mutex = threading.Lock()

def is_model_locked(model_key: str) -> bool:
    with _lock_mutex:
        lock_time = MODELS_LOCK_REGISTRY.get(model_key)
        if lock_time:
            if datetime.now() < lock_time:
                return True
            else:
                del MODELS_LOCK_REGISTRY[model_key]
        return False

def lock_model(model_key: str, seconds: int = 30):
    with _lock_mutex:
        MODELS_LOCK_REGISTRY[model_key] = datetime.now() + timedelta(seconds=seconds)
    print(f"[LOCK] {model_key} hors circuit {seconds}s.")

# ── TIERS ──────────────────────────────────────────────────────────────────────
VALID_TIERS = {"connected_free", "basic", "premium", "ultra", "founder"}

def normalize_tier(raw: str) -> str:
    cleaned = (raw or "").lower().strip()
    if cleaned in VALID_TIERS: return cleaned
    if cleaned == "free":      return "connected_free"
    if "founder" in cleaned:   return "founder"
    if "ultra"   in cleaned:   return "ultra"
    if "premium" in cleaned:   return "premium"
    if "basic"   in cleaned:   return "basic"
    return "connected_free"

def is_paid_tier(tier: str) -> bool:
    return tier in ("basic", "premium", "ultra", "founder")

# ── ERREURS STANDARD ───────────────────────────────────────────────────────────
ERR_QUOTA   = {"action": None, "response": "Le service est temporairement sature. Reessaie dans quelques instants."}
ERR_FINAL   = {"action": None, "response": "Tous les serveurs sont momentanement indisponibles. Reessaie dans 1 minute."}
ERR_CRASH   = {"action": None, "response": "Une erreur inattendue s'est produite. Reessaie dans quelques secondes."}
ERR_HORIZON = {"response": "Le moteur de recherche est temporairement indisponible.", "attributes": [], "matrix": None}

# ── PARSERS JSON ───────────────────────────────────────────────────────────────
def clean_and_parse_json(raw_text):
    if not raw_text or not isinstance(raw_text, str):
        raise ValueError("Reponse vide ou invalide.")
    text = raw_text.strip()
    if text.startswith("```json"): text = text[7:]
    elif text.startswith("```"):   text = text[3:]
    if text.endswith("```"):       text = text[:-3]
    text = text.strip()
    if not text:
        raise ValueError("Reponse vide apres nettoyage.")
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict) and "response" in parsed:
            # Nettoyer les \n littéraux dans la réponse
            if isinstance(parsed["response"], str):
                parsed["response"] = parsed["response"].replace("\\n", "\n").replace("\\\\n", "\n")
            return parsed
    except Exception:
        pass
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group(0))
            if isinstance(parsed, dict) and "response" in parsed:
                if isinstance(parsed["response"], str):
                    parsed["response"] = parsed["response"].replace("\\n", "\n").replace("\\\\n", "\n")
                return parsed
        except Exception:
            pass
    if '"response":' in text:
        res_match = re.search(r'"response"\s*:\s*"([^"]+)"', text)
        if res_match:
            response_text = res_match.group(1).replace("\\n", "\n").replace("\\\\n", "\n")
            return {"action": None, "response": response_text}
    return {"action": None, "response": text}

def clean_and_parse_horizon_json(raw_text):
    if not raw_text or not isinstance(raw_text, str):
        return {"response": "", "attributes": [], "matrix": None}
    text = raw_text.strip()
    if text.startswith("```json"): text = text[7:]
    elif text.startswith("```"):   text = text[3:]
    if text.endswith("```"):       text = text[:-3]
    text = text.strip()
    if not text:
        return {"response": "", "attributes": [], "matrix": None}
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            if isinstance(parsed.get("response"), str):
                parsed["response"] = parsed["response"].replace("\\n", "\n").replace("\\\\n", "\n")
            return parsed
    except Exception:
        pass
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group(0))
            if isinstance(parsed, dict):
                if isinstance(parsed.get("response"), str):
                    parsed["response"] = parsed["response"].replace("\\n", "\n").replace("\\\\n", "\n")
                return parsed
        except Exception:
            pass
    resp_match = re.search(r'"response"\s*:\s*"(.*?)"(?:\s*,|\s*\})', text, re.DOTALL)
    if resp_match:
        response_text = resp_match.group(1).replace('\\n', '\n').replace('\\"', '"').replace("\\\\n", "\n")
        return {"response": response_text, "attributes": [], "matrix": None}
    return {"response": text.replace("**", "").replace("##", "").replace("# ", ""), "attributes": [], "matrix": None}

# ── BUILD CONTENTS ─────────────────────────────────────────────────────────────
def build_gemini_contents(historique_reduit, image_b64, user_message, force_neutral_style):
    contents = []
    for msg in historique_reduit:
        if not isinstance(msg, str) or msg.startswith("__IMAGE__:"):
            continue
        clean_content = msg.split(":", 1)[1].strip() if ":" in msg else msg.strip()
        if "action limit reached" in clean_content.lower() or clean_content == "...":
            continue
        if msg.startswith("Echo:"):
            try:
                parsed = json.loads(clean_content)
                clean_content = parsed.get("response", clean_content)
            except Exception:
                pass
            if force_neutral_style:
                clean_content = "[Analyse technique archivee]"
            contents.append({"role": "model", "parts": [types.Part.from_text(text=clean_content)]})
        elif msg.startswith("You:") or msg.startswith("Toi:"):
            clean_content = re.sub(r'\s{3,}', ' ', clean_content).strip()
            contents.append({"role": "user", "parts": [types.Part.from_text(text=clean_content)]})
    last_parts = []
    if image_b64:
        try:
            header, b64data = image_b64.split(",", 1)
            mime_type = header.split(":")[1].split(";")[0]
            raw_bytes = base64.b64decode(b64data)
            last_parts.append(types.Part.from_bytes(data=raw_bytes, mime_type=mime_type))
        except Exception as e:
            print(f"[WARN] Image error: {e}")
    last_parts.append(types.Part.from_text(text=user_message or "Analyse cette image."))
    contents.append({"role": "user", "parts": last_parts})
    return contents

# ── CONTEXTE PARTAGÉ ───────────────────────────────────────────────────────────
def prepare_shared_context(data, source_override=None):
    user_message     = data.get("message", "")
    calendar_events  = data.get("calendarEvents", {})
    raw_history      = data.get("history", [])
    memory_summary   = data.get("summary", "")
    source           = source_override or (data.get("source") or "chat").lower().strip()
    image_b64        = data.get("image", None)
    selected_buttons = data.get("selectedButtons", [])
    current_expenses = data.get("currentExpenses", [])
    current_calories = data.get("currentCalories", [])
    current_cycle    = data.get("currentCycle", "mois")
    user_tier        = normalize_tier(data.get("userTier", "connected_free"))

    maintenant      = datetime.now()
    date_aujourdhui = maintenant.strftime("%A %d %B %Y")
    annee_en_cours  = maintenant.strftime("%Y")

    filtered_calendar = {}
    date_debut = (maintenant - timedelta(days=31)).date()
    date_fin   = maintenant.date()
    try:
        for date_str, events in calendar_events.items():
            if re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
                date_evt = datetime.strptime(date_str, "%Y-%m-%d").date()
                if date_debut <= date_evt <= date_fin:
                    filtered_calendar[date_str] = events
            else:
                filtered_calendar[date_str] = events
    except Exception as e:
        print(f"[WARN] Calendrier: {e}")
        filtered_calendar = calendar_events

    base_system_prompt = generate_system_prompt(
        source=source,
        selected_buttons=selected_buttons,
        date_aujourdhui=date_aujourdhui,
        annee_en_cours=annee_en_cours,
        user_tier=user_tier,
        filtered_calendar=filtered_calendar,
        current_expenses=current_expenses,
        current_calories=current_calories,
        current_cycle=current_cycle,
    )

    system_prompt = base_system_prompt
    if memory_summary and user_tier not in ("connected_free", "basic"):
        system_prompt += f"\n\nLONG TERM MEMORY\n================\n{memory_summary}\n"
    system_prompt += "\n\nCRITICAL SAFETY DIRECTIVE: Only trigger actions explicitly demanded in the LATEST message."

    taille_memoire    = 30 if user_tier in ("ultra", "founder") else (15 if user_tier in ("basic", "premium") else 12)
    output_tokens     = PAID_MAX_TOKENS if is_paid_tier(user_tier) else FREE_MAX_TOKENS
    historique_ajuste = raw_history[-taille_memoire:]
    force_neutral     = len(selected_buttons) > 0 or source == "vitality"
    gemini_contents   = build_gemini_contents(historique_ajuste, image_b64, user_message, force_neutral)

    messages_openai = [{"role": "system", "content": system_prompt}]
    for msg in historique_ajuste:
        if not isinstance(msg, str) or msg.startswith("__IMAGE__:"):
            continue
        clean_content = msg.split(":", 1)[1].strip() if ":" in msg else msg.strip()
        if "action limit reached" in clean_content.lower() or clean_content == "...":
            continue
        if msg.startswith("You:") or msg.startswith("Toi:"):
            messages_openai.append({"role": "user", "content": re.sub(r'\s{3,}', ' ', clean_content).strip()})
        elif msg.startswith("Echo:"):
            try:
                parsed = json.loads(clean_content)
                clean_content = parsed.get("response", clean_content)
            except Exception:
                pass
            messages_openai.append({"role": "assistant", "content": clean_content})
    if user_message:
        messages_openai.append({"role": "user", "content": user_message})

    return {
        "system_prompt":   system_prompt,
        "output_tokens":   output_tokens,
        "gemini_contents": gemini_contents,
        "messages_openai": messages_openai,
        "user_tier":       user_tier,
    }

# ── CALL HELPERS ───────────────────────────────────────────────────────────────
def call_with_timeout(fn, timeout_sec):
    result = [None]
    error  = [None]
    def target():
        try:
            result[0] = fn()
        except Exception as e:
            error[0] = e
    t = threading.Thread(target=target, daemon=True)
    t.start()
    t.join(timeout=timeout_sec)
    if t.is_alive():
        raise TimeoutError(f"Timeout apres {timeout_sec}s")
    if error[0]:
        raise error[0]
    return result[0]

def call_gemini(client, model_key, ctx, timeout=25, temperature=0.5):
    def fn():
        return client.models.generate_content(
            model=MODELS[model_key],
            contents=ctx["gemini_contents"],
            config=types.GenerateContentConfig(
                system_instruction=ctx["system_prompt"],
                max_output_tokens=ctx["output_tokens"],
                temperature=temperature,
            )
        )
    return call_with_timeout(fn, timeout)

def call_gemini_with_search(client, model_key, ctx, timeout=30, temperature=0.5):
    def fn():
        return client.models.generate_content(
            model=MODELS[model_key],
            contents=ctx["gemini_contents"],
            config=types.GenerateContentConfig(
                system_instruction=ctx["system_prompt"],
                max_output_tokens=ctx["output_tokens"],
                temperature=temperature,
                tools=[types.Tool(google_search=types.GoogleSearch())],
            )
        )
    return call_with_timeout(fn, timeout)

def call_openrouter(model_key, ctx, temp=0.5, timeout=20.0):
    if client_openrouter is None:
        raise RuntimeError("OpenRouter non configuré")
    res = client_openrouter.chat.completions.create(
        model=MODELS[model_key],
        messages=ctx["messages_openai"],
        temperature=temp,
        max_tokens=ctx.get("output_tokens", FREE_MAX_TOKENS),
        timeout=timeout,
    )
    content = res.choices[0].message.content
    if content is None:
        raise ValueError(f"Reponse vide de {model_key}")
    return content

def call_deepseek(ctx, temp=0.5, timeout=20.0):
    if client_deepseek is None:
        raise RuntimeError("DeepSeek non configuré")
    res = client_deepseek.chat.completions.create(
        model=MODELS["deepseek"],
        messages=ctx["messages_openai"],
        temperature=temp,
        max_tokens=ctx.get("output_tokens", FREE_MAX_TOKENS),
        timeout=timeout,
    )
    content = res.choices[0].message.content
    if content is None:
        raise ValueError("Reponse vide de deepseek")
    return content

def call_glm(ctx, temp=0.5, timeout=20.0):
    if client_zai is None:
        raise RuntimeError("Z.AI non configuré")
    res = client_zai.chat.completions.create(
        model=MODELS["glm"],
        messages=ctx["messages_openai"],
        temperature=temp,
        max_tokens=ctx.get("output_tokens", FREE_MAX_TOKENS),
        timeout=timeout,
    )
    content = res.choices[0].message.content
    if content is None:
        raise ValueError("Reponse vide de glm")
    return content

# ── DISPATCH COMMUN POUR LES CASCADES ──────────────────────────────────────────
def _dispatch_step(provider, model_key, ctx, timeout, parser):
    """Exécute une étape de cascade et retourne le résultat parsé."""
    if provider == "g":
        r = call_gemini(client_gemini_paid, model_key, ctx, timeout=timeout, temperature=0.5)
        return parser(r.text)
    elif provider == "gs":
        r = call_gemini_with_search(client_gemini_paid, model_key, ctx, timeout=timeout, temperature=0.5)
        return parser(r.text)
    elif provider == "ds":
        r = call_deepseek(ctx, temp=0.5, timeout=float(timeout))
        return parser(r)
    elif provider == "z":
        r = call_glm(ctx, temp=0.5, timeout=float(timeout))
        return parser(r)
    else:  # "or"
        r = call_openrouter(model_key, ctx, temp=0.5, timeout=float(timeout))
        return parser(r)

# ── CASCADE PAYANTE PAR TIER (NE COUVRE PAS /vitality) ─────────────────────────
# Utilisée par /home, /chat, /history (payant). Cascade différente selon basic/premium/ultra/founder.
def run_paid_cascade(ctx, page_timeout: int = 10):
    tier = ctx["user_tier"]
    t    = page_timeout

    if tier == "basic":
        steps = [
            ("ds", "deepseek", t),
            ("or", "llama",    t),
            ("or", "hy3",      t),
        ]
    elif tier == "premium":
        steps = [
            ("ds", "deepseek", t),
            ("or", "llama",    t),
            ("ds", "deepseek", t),
        ]
    elif tier == "ultra":
        steps = [
            ("ds", "deepseek",          t),
            ("g",  "gemini_paid_ultra", t),
            ("ds", "deepseek",          t),
        ]
    else:  # founder
        steps = [
            ("g",  "gemini_paid_founder", t),
            ("g",  "gemini_paid_ultra",   t),
            ("ds", "deepseek",            t),
        ]

    for provider, model_key, timeout in steps:
        if is_model_locked(model_key):
            continue
        try:
            return _dispatch_step(provider, model_key, ctx, timeout, clean_and_parse_json)
        except Exception as e:
            print(f"[PAID {tier}] Echec {model_key} ({e})")
            lock_model(model_key, 30)
    return ERR_FINAL

# ── CASCADE FREE ───────────────────────────────────────────────────────────────
def run_free_cascade(steps, ctx, parser=None, is_horizon=False):
    if parser is None:
        parser = clean_and_parse_json
    if not is_horizon and get_failover_count() >= MAX_FREE_FAILOVERS:
        return ERR_QUOTA
    for i, (provider, model_key, timeout) in enumerate(steps):
        if provider == "or" and client_openrouter is None:
            continue
        if provider == "ds" and client_deepseek is None:
            continue
        if provider == "z" and client_zai is None:
            continue
        if is_model_locked(model_key):
            continue
        is_last = (i == len(steps) - 1)
        try:
            result = _dispatch_step(provider, model_key, ctx, timeout, parser)
            if is_last and not is_horizon:
                increment_failover_count()
            return result
        except Exception as e:
            print(f"[FREE] Echec {model_key} ({e})")
            lock_model(model_key, 30)
    return ERR_FINAL

# ── ROUTES ─────────────────────────────────────────────────────────────────────

@app.route("/ping")
def ping():
    return jsonify({"status": "awake"})

# ── /horizon-warmup ───────────────────────────────────────────────────────────
HORIZON_WARMUP_CACHE: dict = {}

@app.route("/horizon-warmup", methods=["POST"])
def horizon_warmup():
    try:
        data    = request.json or {}
        partial = (data.get("partial") or "").strip()
        if not partial or len(partial) < 3:
            return jsonify({"status": "too_short"}), 200
        prompt = (
            f"Analyse cette intention de recherche : \"{partial}\"\n\n"
            f"Réponds en JSON avec exactement ces clés :\n"
            f"{{\"intent\": \"[recherche_locale|prix|comparaison|definition|actualite|autre]\","
            f" \"keywords\": [\"mot1\", \"mot2\"], \"lang\": \"fr|en\"}}\n"
            f"JSON uniquement, rien d'autre."
        )
        ctx = {
            "system_prompt": "Tu es un classificateur d'intentions de recherche. Réponds uniquement en JSON valide.",
            "output_tokens": 150,
            "messages_openai": [
                {"role": "system", "content": "Tu es un classificateur d'intentions de recherche. Réponds uniquement en JSON valide."},
                {"role": "user",   "content": prompt},
            ],
            "gemini_contents": [{"role": "user", "parts": [types.Part.from_text(text=prompt)]}],
        }
        try:
            raw    = call_openrouter("ling", ctx, temp=0.1, timeout=4.0)
            parsed = clean_and_parse_json(raw)
            HORIZON_WARMUP_CACHE[partial[:50]] = parsed
            print(f"[WARMUP] Intent détecté : {parsed.get('response', parsed)}")
            return jsonify({"status": "ready", "intent": parsed})
        except Exception as e:
            print(f"[WARMUP] Ling échec ({e})")
            return jsonify({"status": "fallback"})
    except Exception as e:
        print(f"[WARMUP] Erreur : {e}")
        return jsonify({"status": "error"})

# ── /home ──────────────────────────────────────────────────────────────────────
# FREE : DeepSeek → tencent (hy3) → Llama
@app.route("/home", methods=["POST"])
def home():
    try:
        data = request.json or {}
        ctx  = prepare_shared_context(data, source_override="home")
        if is_paid_tier(ctx["user_tier"]):
            return jsonify(run_paid_cascade(ctx, page_timeout=10))
        steps = [
            ("ds", "deepseek",  8),
            ("or", "hy3",      12),
            ("or", "llama",    15),
        ]
        return jsonify(run_free_cascade(steps, ctx))
    except Exception as e:
        print(f"Erreur /home: {e}")
        return jsonify(ERR_CRASH), 500

# ── /chat ──────────────────────────────────────────────────────────────────────
# FREE : DeepSeek → tencent (hy3) → Llama
@app.route("/chat", methods=["POST"])
def chat():
    try:
        data = request.json or {}
        ctx  = prepare_shared_context(data, source_override="chat")
        if is_paid_tier(ctx["user_tier"]):
            return jsonify(run_paid_cascade(ctx, page_timeout=10))
        steps = [
            ("ds", "deepseek",  8),
            ("or", "hy3",      12),
            ("or", "llama",    15),
        ]
        return jsonify(run_free_cascade(steps, ctx))
    except Exception as e:
        print(f"Erreur /chat: {e}")
        return jsonify(ERR_CRASH), 500

# ── /vitality ─────────────────────────────────────────────────────────────────
# INTACT : seule route où GLM → DeepSeek → Llama reste en place, FREE et PAYANT.
def _run_vitality_paid_cascade(ctx, page_timeout: int = 20):
    tier = ctx["user_tier"]
    steps = [
        ("z",  "glm",      15),
        ("ds", "deepseek", page_timeout),
        ("or", "llama",    page_timeout),
    ]
    for provider, model_key, timeout in steps:
        if is_model_locked(model_key):
            continue
        try:
            return _dispatch_step(provider, model_key, ctx, timeout, clean_and_parse_json)
        except Exception as e:
            print(f"[PAID VITALITY {tier}] Echec {model_key} ({e})")
            lock_model(model_key, 30)
    return ERR_FINAL

@app.route("/vitality", methods=["POST"])
def vitality():
    try:
        data = request.json or {}
        ctx  = prepare_shared_context(data, source_override="vitality")
        steps = [
            ("z",  "glm",      15),
            ("ds", "deepseek", 20),
            ("or", "llama",    20),
        ]
        if is_paid_tier(ctx["user_tier"]):
            result = _run_vitality_paid_cascade(ctx, page_timeout=20)
        else:
            result = run_free_cascade(steps, ctx)
        print(f"[VITALITY] action retournée: {result.get('action')}")
        return jsonify(result)
    except Exception as e:
        print(f"Erreur /vitality: {e}")
        return jsonify(ERR_CRASH), 500

# ── /history ──────────────────────────────────────────────────────────────────
# FREE fallback (après owl) : DeepSeek garde sa place, GLM passe en dernier
@app.route("/history", methods=["POST"])
def history():
    try:
        data = request.json or {}
        ctx  = prepare_shared_context(data, source_override="history")
        if is_paid_tier(ctx["user_tier"]):
            return jsonify(run_paid_cascade(ctx, page_timeout=12))
        if get_failover_count() >= MAX_FREE_FAILOVERS:
            return jsonify(ERR_QUOTA)
        result = None
        if not is_model_locked("owl"):
            try:
                r      = call_openrouter("owl", ctx, temp=0.5, timeout=15.0)
                result = clean_and_parse_json(r)
                print("[HISTORY] owl-alpha OK")
            except Exception as e:
                print(f"[HISTORY] owl-alpha echec ({e})")
            finally:
                lock_model("owl", 60)
        if result is None:
            fallback_steps = [
                ("or", "mistral",  20),
                ("ds", "deepseek", 20),
            ]
            result = run_free_cascade(fallback_steps, ctx)
        return jsonify(result)
    except Exception as e:
        print(f"Erreur /history: {e}")
        return jsonify(ERR_CRASH), 500

# ── /books ────────────────────────────────────────────────────────────────────
# Tous tiers : DeepSeek → Gemini 3.1 Flash Lite → DeepSeek (filet) — INTACT, pas de GLM ici
@app.route("/books", methods=["POST"])
def books():
    try:
        data       = request.json or {}
        message    = (data.get("message") or "").strip()
        history    = data.get("history", [])
        tier       = normalize_tier(data.get("userTier", "connected_free"))
        buttons    = data.get("selectedButtons", [])
        book_title = (data.get("bookTitle") or "").strip()
        output_tokens = 2048 if is_paid_tier(tier) else 1024

        INJECT_KEYWORDS = ["inject", "injecte", "insere", "ecris ici", "write here", "add this"]
        wants_inject    = any(kw in message.lower() for kw in INJECT_KEYWORDS)

        mode_prompts = {
            "creative": "Tu es en mode Creatif. Genere du contenu litteraire original avec un style soigne.",
            "ideas":    "Tu es en mode Idees. Propose des pistes narratives et rebondissements.",
            "critical": "Tu es en mode Critique. Analyse le texte : rythme, coherence, clarte.",
        }
        active_mode      = buttons[0] if buttons else None
        mode_instruction = mode_prompts.get(active_mode, "Tu es un assistant d'ecriture creatif et polyvalent.")
        inject_instruction = ""
        if wants_inject:
            inject_instruction = (
                "\n\nL'utilisateur veut que tu INJECTES du texte dans son livre. "
                "Genere le passage et termine avec :\n<<<INJECT_TEXT>>>\n[texte propre sans HTML]\n<<<END_INJECT>>>"
            )
        system_prompt = f"{mode_instruction}{inject_instruction}\n\nLivre : \"{book_title}\". Reponds en moins de 400 mots. Sois direct et elegant."

        def build_openai_messages(hist):
            msgs = [{"role": "system", "content": system_prompt}]
            for msg in hist[-10:]:
                if not isinstance(msg, str): continue
                if msg.startswith("You: ") or msg.startswith("Toi: "):
                    msgs.append({"role": "user", "content": msg.split(":", 1)[1].strip()})
                elif msg.startswith("Echo: "):
                    msgs.append({"role": "assistant", "content": msg.split(":", 1)[1].strip()})
            msgs.append({"role": "user", "content": message})
            return msgs

        def run_books_call_deepseek(timeout):
            if client_deepseek is None:
                raise RuntimeError("DeepSeek non configuré")
            res = client_deepseek.chat.completions.create(
                model=MODELS["deepseek"],
                messages=build_openai_messages(history),
                temperature=0.5,
                max_tokens=output_tokens,
                timeout=float(timeout),
            )
            content = res.choices[0].message.content
            if not content:
                raise ValueError("Reponse vide de deepseek")
            return content

        def run_books_call_gemini_flash(timeout):
            gemini_hist = []
            for msg in history[-10:]:
                if msg.startswith("You: "):
                    gemini_hist.append(types.Content(role="user",  parts=[types.Part.from_text(text=msg[5:])]))
                elif msg.startswith("Echo: "):
                    if gemini_hist:
                        gemini_hist.append(types.Content(role="model", parts=[types.Part.from_text(text=msg[6:])]))
            gemini_hist.append(types.Content(role="user", parts=[types.Part.from_text(text=message)]))
            def fn():
                return client_gemini_paid.models.generate_content(
                    model=MODELS["gemini_paid_ultra"],
                    contents=gemini_hist,
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        max_output_tokens=output_tokens,
                        temperature=0.5,
                    )
                )
            resp = call_with_timeout(fn, timeout)
            if not resp or not resp.text:
                raise ValueError("Reponse Gemini vide")
            return resp.text

        books_steps = [
            ("ds_books",     20),
            ("gemini_flash", 20),
            ("ds_books",     25),
        ]

        full_text = ""
        for provider, timeout in books_steps:
            try:
                if provider == "ds_books":
                    if is_model_locked("deepseek"):
                        continue
                    full_text = run_books_call_deepseek(timeout)
                elif provider == "gemini_flash":
                    if is_model_locked("gemini_paid_ultra"):
                        continue
                    full_text = run_books_call_gemini_flash(timeout)
                if full_text:
                    break
            except Exception as e:
                model_tag = "deepseek" if provider == "ds_books" else "gemini_paid_ultra"
                print(f"[BOOKS] Echec {model_tag} ({e})")
                lock_model(model_tag, 30)

        if not full_text:
            return jsonify({"response": ERR_FINAL["response"], "inject": False, "action": None})

        if wants_inject and "<<<INJECT_TEXT>>>" in full_text:
            parts      = full_text.split("<<<INJECT_TEXT>>>")
            response   = parts[0].strip()
            inject_raw = parts[1].split("<<<END_INJECT>>>")[0].strip() if len(parts) > 1 else ""
            return jsonify({"response": response or "Voici le passage.", "inject": True, "inject_text": inject_raw, "action": None})
        return jsonify({"response": full_text, "inject": False, "action": None})

    except Exception as e:
        print(f"Erreur /books: {e}")
        return jsonify({"response": ERR_CRASH["response"], "inject": False, "action": None}), 500

# ── /horizon ──────────────────────────────────────────────────────────────────
# INTACT : pas de GLM dans cette route
@app.route("/horizon", methods=["POST"])
def horizon():
    try:
        data  = request.json or {}
        query = (data.get("query") or "").strip()
        if not query:
            return jsonify({"error": "L'intention d'exploration est vide."}), 400
        if data.get("warmup") or query == "...":
            return jsonify({"status": "warm"}), 200

        maintenant = datetime.now()
        date_str   = maintenant.strftime("%A %d %B %Y")
        heure_str  = maintenant.strftime("%H:%M")
        user_tier  = normalize_tier(data.get("userTier", "connected_free"))

        horizon_message = (
            f"DATE ET HEURE ACTUELLES : {date_str}, {heure_str} (heure locale du serveur).\n"
            f"ANNEE EN COURS : 2026. Toutes les informations doivent etre actuelles et valides en 2026.\n\n"
            f"INSTRUCTION OBLIGATOIRE — GOOGLE SEARCH ACTIVE :\n"
            f"Tu disposes d'un outil de recherche Google en temps reel. "
            f"Tu DOIS l'utiliser pour repondre a cette requete. "
            f"N'utilise PAS ta memoire interne comme source principale. "
            f"Cherche d'abord. Transmets ce que tu trouves. Rien d'autre.\n\n"
            f"REQUETE : {query}\n\n"
            f"REGLES DE TRANSMISSION :\n"
            f"- Retourne uniquement des informations confirmees par ta recherche Google.\n"
            f"- Si une donnee (adresse, telephone, horaire, prix, URL) est introuvable : utilise le jeton approprie.\n"
            f"- Retourne jusqu'a 10 resultats confirmes — moins si tu n'en trouves pas davantage.\n"
            f"- Une reponse incomplete est acceptable. Une reponse inventee est un echec.\n\n"
            f"Reponds UNIQUEMENT en JSON valide avec les cles : response, attributes, matrix."
        )

        horizon_system = generate_system_prompt(
            source="horizonweb",
            selected_buttons=[],
            date_aujourdhui=date_str,
            annee_en_cours="2026",
            user_tier=user_tier,
            filtered_calendar={},
        )

        ctx = {
            "system_prompt":   horizon_system,
            "output_tokens":   PAID_MAX_TOKENS if is_paid_tier(user_tier) else FREE_MAX_TOKENS,
            "gemini_contents": [{"role": "user", "parts": [types.Part.from_text(text=horizon_message)]}],
            "messages_openai": [
                {"role": "system", "content": horizon_system},
                {"role": "user",   "content": horizon_message},
            ],
            "user_tier": user_tier,
        }

        parsed = None
        horizon_steps = [
            ("gs", "gemini_paid_standard", 8),
            ("gs", "gemini_paid_ultra",    8),
            ("ds", "deepseek",             8),
        ]

        for attempt in range(3):
            for provider, model_key, timeout in horizon_steps:
                try:
                    print(f"[HORIZON] Pass {attempt+1} — {model_key}")
                    if provider == "gs":
                        r      = call_gemini_with_search(client_gemini_paid, model_key, ctx, timeout=timeout, temperature=0.5)
                        parsed = clean_and_parse_horizon_json(r.text)
                    else:
                        r      = call_deepseek(ctx, temp=0.5, timeout=float(timeout))
                        parsed = clean_and_parse_horizon_json(r)
                    if parsed.get("response", "").strip():
                        print(f"[HORIZON] Succes — {model_key}")
                        break
                    else:
                        print(f"[HORIZON] Reponse vide — {model_key}")
                except Exception as e:
                    print(f"[HORIZON] Echec {model_key} ({e})")
            if parsed and parsed.get("response", "").strip():
                break
            if attempt < 2:
                print(f"[HORIZON] Retry complet pass {attempt+2}...")

        if not parsed or not parsed.get("response", "").strip():
            return jsonify(ERR_HORIZON)

        if not isinstance(parsed.get("matrix"), dict):
            parsed["matrix"] = None
        if not isinstance(parsed.get("attributes"), list):
            parsed["attributes"] = []

        return jsonify(parsed)

    except Exception as e:
        print(f"[HORIZON] Erreur critique: {e}")
        return jsonify(ERR_HORIZON), 500

# ── /memory-summary ───────────────────────────────────────────────────────────
# INTACT : pas de GLM dans cette route
@app.route("/memory-summary", methods=["POST"])
def memory_summary():
    try:
        data         = request.json or {}
        messages     = data.get("messages", [])
        prev_summary = data.get("summary", "")
        user_tier    = normalize_tier(data.get("userTier", "connected_free"))

        if not messages:
            return jsonify({"summary": prev_summary or ""})

        history_text = "\n".join(messages[-30:])
        prompt = (
            "Tu es un système de compression de mémoire conversationnelle.\n"
            "Ton rôle : produire un résumé dense et factuel de la conversation ci-dessous.\n"
            "Ce résumé sera injecté comme mémoire long terme dans les prochaines conversations.\n\n"
            f"RÉSUMÉ PRÉCÉDENT :\n{prev_summary or 'Aucun'}\n\n"
            f"NOUVEAUX MESSAGES :\n{history_text}\n\n"
            "INSTRUCTIONS :\n"
            "- Garde les faits importants, préférences, décisions, contexte clé.\n"
            "- Ignore les salutations et messages vides.\n"
            "- Sois concis : max 300 mots.\n"
            "- Réponds UNIQUEMENT avec le résumé, sans introduction ni explication."
        )

        ctx = {
            "system_prompt":   "Tu es un compresseur de mémoire conversationnelle. Réponds uniquement avec le résumé demandé.",
            "output_tokens":   600,
            "gemini_contents": [{"role": "user", "parts": [types.Part.from_text(text=prompt)]}],
            "messages_openai": [
                {"role": "system", "content": "Tu es un compresseur de mémoire conversationnelle."},
                {"role": "user",   "content": prompt},
            ],
            "user_tier": user_tier,
        }

        try:
            r       = call_openrouter("ling", ctx, temp=0.1, timeout=15.0)
            summary = r.strip() if r else ""
            if summary:
                return jsonify({"summary": summary})
        except Exception as e:
            print(f"[MEMORY] ling echec ({e}), fallback nova-lite")

        try:
            r       = call_openrouter("mistral", ctx, temp=0.1, timeout=15.0)
            summary = r.strip() if r else ""
            if summary:
                return jsonify({"summary": summary})
        except Exception as e:
            print(f"[MEMORY] Echec total ({e})")

        return jsonify({"summary": prev_summary or ""})

    except Exception as e:
        print(f"[MEMORY] Erreur critique: {e}")
        return jsonify({"summary": ""}), 500

# ── /export ───────────────────────────────────────────────────────────────────
def px_to_pt(px_str: str) -> float:
    try:
        return round(float(str(px_str).replace("px","").strip()) * 0.75, 1)
    except Exception:
        return 11.0

def parse_style(style_str: str) -> dict:
    result = {}
    for part in (style_str or "").split(";"):
        part = part.strip()
        if ":" in part:
            k, v = part.split(":", 1)
            result[k.strip().lower()] = v.strip()
    return result

def html_to_docx(html: str, title: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup, NavigableString, Tag

    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run  = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    doc.add_paragraph()

    soup = BeautifulSoup(html, "html.parser")

    global_font_size_pt = 11.0
    global_font_family  = "Georgia"
    wrapper = soup.find("div", style=True)
    if wrapper:
        ws = parse_style(wrapper.get("style", ""))
        if "font-size" in ws:
            global_font_size_pt = px_to_pt(ws["font-size"])
        if "font-family" in ws:
            global_font_family = ws["font-family"].split(",")[0].strip().strip('"\'')

    ALIGN_MAP = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def add_run_from_node(para, node, inherited: dict):
        if isinstance(node, NavigableString):
            text = str(node)
            if not text.strip() and text != " ":
                return
            run = para.add_run(text)
            run.font.size = Pt(inherited.get("font_size_pt", global_font_size_pt))
            run.font.name = inherited.get("font_family", global_font_family)
            run.bold      = inherited.get("bold", False)
            run.italic    = inherited.get("italic", False)
            return
        if not isinstance(node, Tag):
            return
        tag  = node.name.lower() if node.name else ""
        styl = parse_style(node.get("style", ""))
        ctx  = dict(inherited)
        if "font-size"   in styl: ctx["font_size_pt"] = px_to_pt(styl["font-size"])
        if "font-family" in styl: ctx["font_family"]  = styl["font-family"].split(",")[0].strip().strip('"\'')
        if "font-weight" in styl and styl["font-weight"] in ("bold","700","800","900"): ctx["bold"]   = True
        if "font-style"  in styl and styl["font-style"] == "italic":                   ctx["italic"] = True
        if tag in ("strong", "b"): ctx["bold"]   = True
        if tag in ("em", "i"):     ctx["italic"] = True
        if tag == "br":
            para.add_run().add_break(); return
        for child in node.children:
            add_run_from_node(para, child, ctx)

    def process_block(node):
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                p   = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(global_font_size_pt)
                run.font.name = global_font_family
            return
        if not isinstance(node, Tag):
            return
        tag  = node.name.lower() if node.name else ""
        styl = parse_style(node.get("style", ""))
        if node.get("data-page-break"):
            doc.add_page_break(); return
        if tag in ("script", "style", "head"):
            return
        if tag in ("div", "section", "article", "main", "body", "html"):
            for child in node.children:
                process_block(child)
            return
        heading_sizes = {"h1": 24.0, "h2": 18.0, "h3": 14.0}
        if tag in heading_sizes:
            para = doc.add_paragraph()
            ctx  = {"font_size_pt": heading_sizes[tag], "bold": True, "font_family": global_font_family}
            if "text-align" in styl:
                para.alignment = ALIGN_MAP.get(styl["text-align"], WD_ALIGN_PARAGRAPH.LEFT)
            for child in node.children:
                add_run_from_node(para, child, ctx)
            return
        if tag in ("p", "li"):
            para = doc.add_paragraph()
            ctx  = {"font_size_pt": global_font_size_pt, "font_family": global_font_family}
            if "text-align" in styl:
                para.alignment = ALIGN_MAP.get(styl["text-align"], WD_ALIGN_PARAGRAPH.LEFT)
            if "font-size"  in styl:
                ctx["font_size_pt"] = px_to_pt(styl["font-size"])
            for child in node.children:
                add_run_from_node(para, child, ctx)
            return
        if tag in ("span", "strong", "em", "b", "i", "u"):
            para = doc.add_paragraph()
            ctx  = {"font_size_pt": global_font_size_pt, "font_family": global_font_family}
            add_run_from_node(para, node, ctx)
            return
        for child in node.children:
            process_block(child)

    root = wrapper if wrapper else soup
    for child in root.children:
        process_block(child)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/export", methods=["POST"])
def export_route():
    data  = request.get_json(silent=True) or {}
    fmt   = (data.get("format") or "").lower().strip()
    title = (data.get("title")  or "Document Echo AI").strip()
    html  = (data.get("html")   or "").strip()
    if not html:
        return jsonify({"error": "Contenu vide."}), 400
    safe = "".join(c for c in title if c.isalnum() or c in " _-").strip().replace(" ", "_") or "document"
    try:
        if fmt == "txt":
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            txt  = soup.get_text(separator="\n").replace('&nbsp;', ' ')
            return send_file(io.BytesIO(txt.encode("utf-8")), mimetype="text/plain", as_attachment=True, download_name=f"{safe}.txt")

        elif fmt == "pdf":
            try:
                from xhtml2pdf import pisa
                from bs4 import BeautifulSoup
                soup    = BeautifulSoup(html, "html.parser")
                wrapper = soup.find("div", style=True)
                fs_pt   = "11pt"
                ff      = "Georgia, serif"
                if wrapper:
                    ws = parse_style(wrapper.get("style", ""))
                    if "font-size"   in ws: fs_pt = f"{px_to_pt(ws['font-size'])}pt"
                    if "font-family" in ws: ff    = ws["font-family"]
                styled = (
                    f'<html><head><meta charset="utf-8">'
                    f'<style>body{{font-family:{ff};font-size:{fs_pt};line-height:1.6;color:#18181b}}'
                    f'h1{{font-size:24pt}}h2{{font-size:18pt}}h3{{font-size:14pt}}'
                    f'p{{margin-bottom:10px}}</style></head>'
                    f'<body><h1>{title}</h1>{html}</body></html>'
                )
                buf = io.BytesIO()
                pisa.CreatePDF(io.StringIO(styled), dest=buf)
                buf.seek(0)
                return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=f"{safe}.pdf")
            except ImportError:
                return jsonify({"error": "xhtml2pdf non installe."}), 503

        elif fmt == "docx":
            try:
                buf = html_to_docx(html, title)
                return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=f"{safe}.docx")
            except ImportError:
                return jsonify({"error": "python-docx ou beautifulsoup4 non installe."}), 503

        elif fmt == "epub":
            try:
                from ebooklib import epub
                book = epub.EpubBook()
                book.set_identifier(f"echo-{int(datetime.now().timestamp())}")
                book.set_title(title); book.set_language("fr"); book.add_author("Echo AI")
                ch = epub.EpubHtml(title=title, file_name="ch1.xhtml", lang="fr")
                ch.content = f"<html><body><h1>{title}</h1>{html}</body></html>"
                book.add_item(ch)
                book.toc   = (epub.Link("ch1.xhtml", title, "ch1"),)
                book.spine = ["nav", ch]
                book.add_item(epub.EpubNav())
                book.add_item(epub.EpubNcx())
                buf = io.BytesIO(); epub.write_epub(buf, book, {}); buf.seek(0)
                return send_file(buf, mimetype="application/epub+zip", as_attachment=True, download_name=f"{safe}.epub")
            except ImportError:
                return jsonify({"error": "EbookLib non installe."}), 503

        else:
            return jsonify({"error": f"Format '{fmt}' non supporte."}), 400

    except Exception as e:
        print(f"[EXPORT ERROR] {fmt}: {e}")
        return jsonify({"error": str(e)}), 500

# ── /api/analyse-avis — Scanner Sonar (Perplexity) via OpenRouter ───────────
@app.route("/api/analyse-avis", methods=["POST"])
def analyse_avis():
    """
    Route pour analyser les avis d'un produit via Perplexity Sonar (web search natif).
    Accepte : {"url": "https://amazon.ca/..."}
    Retourne : {"product_name": "...", "positives": [...], "negatives": [...]}
    """
    try:
        data = request.json or {}
        url = (data.get("url") or "").strip()
        print(f"[ANALYSE] URL reçue: {url[:80]}")

        if not url:
            return jsonify({"error": "URL requise"}), 400

        if client_openrouter is None:
            print("[ANALYSE] ERREUR: client_openrouter est None")
            return jsonify({"error": "OpenRouter non configure"}), 503

        print(f"[ANALYSE] OpenRouter OK, DeepSeek: {client_deepseek is not None}")

        # Extraire l'ASIN si URL Amazon
        asin_match = re.search(r"/dp/([A-Z0-9]{10})", url)
        asin_hint = f" (ASIN: {asin_match.group(1)})" if asin_match else ""
        print(f"[ANALYSE] ASIN: {asin_hint or 'non detecte'}")

        system_prompt = (
            "Tu es un expert en analyse produit e-commerce. "
            "A partir d'une URL ou d'un ASIN Amazon, identifie le produit "
            "et fournis une analyse honnete basee sur ta connaissance des avis clients. "
            "Sois factuel, chirurgical, anti-bullshit. "
            "Reponds UNIQUEMENT avec ce JSON : "
            "{\"product_name\": \"nom complet\", \"positives\": [\"p1\",\"p2\",\"p3\",\"p4\",\"p5\"], \"negatives\": [\"n1\",\"n2\",\"n3\",\"n4\",\"n5\"]}"
        )

        user_prompt = (
            f"Analyse ce produit{asin_hint} : {url}\n\n"
            f"Donne EXACTEMENT 5 points forts et 5 defauts reels "
            f"bases sur les avis clients connus. Chaque point max 12 mots. "
            f"JSON uniquement."
        )

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        print(f"[ANALYSE] gpt-4o-mini-search-preview : {url}")
        raw_response = None

        # Tentative 1 : GPT-4o Mini Search (web search natif, cheap)
        if client_openrouter is not None:
            try:
                res = client_openrouter.chat.completions.create(
                    model="openai/gpt-4o-mini-search-preview",
                    messages=messages,
                    temperature=0.3,
                    max_tokens=700,
                    timeout=30.0,
                )
                raw_response = res.choices[0].message.content
                print("[ANALYSE] GPT-4o-mini-search OK")
            except Exception as e:
                print(f"[ANALYSE] GPT-4o-mini-search echec ({e}), fallback DeepSeek...")

        # Tentative 2 : DeepSeek (connaissance training)
        if not raw_response and client_deepseek is not None:
            try:
                res = client_deepseek.chat.completions.create(
                    model=MODELS["deepseek"],
                    messages=messages,
                    temperature=0.3,
                    max_tokens=700,
                    timeout=20.0,
                )
                raw_response = res.choices[0].message.content
                print("[ANALYSE] DeepSeek fallback OK")
            except Exception as e:
                print(f"[ANALYSE] DeepSeek echec ({e})")

        # Tentative 3 : Llama
        if not raw_response and client_openrouter is not None:
            try:
                res = client_openrouter.chat.completions.create(
                    model=MODELS["llama"],
                    messages=messages,
                    temperature=0.3,
                    max_tokens=700,
                    timeout=20.0,
                )
                raw_response = res.choices[0].message.content
                print("[ANALYSE] Llama fallback OK")
            except Exception as e:
                print(f"[ANALYSE] Llama echec ({e})")

        if not raw_response:
            return jsonify({"error": "Analyse indisponible, reessaie dans quelques secondes"}), 503

        # Parser la reponse
        try:
            parsed = clean_and_parse_json(raw_response)

            # Sonar peut repondre en texte libre — extraire depuis "response" si besoin
            positives = parsed.get("positives")
            negatives = parsed.get("negatives")
            product_name = parsed.get("product_name") or "Produit Analyse"

            if not isinstance(positives, list) or len(positives) == 0:
                positives = ["Donnees insuffisantes"]
            if not isinstance(negatives, list) or len(negatives) == 0:
                negatives = ["Donnees insuffisantes"]

            return jsonify({
                "product_name": product_name,
                "positives": positives[:5],
                "negatives": negatives[:5],
            })

        except Exception as e:
            print(f"[SONAR] Erreur parsing: {e} — raw: {raw_response[:200]}")
            return jsonify({"error": "Erreur parsing reponse"}), 500

    except Exception as e:
        import traceback
        print(f"[ANALYSE] Erreur critique: {e}")
        print(f"[ANALYSE] Traceback: {traceback.format_exc()}")
        return jsonify({"error": "Erreur serveur interne"}), 500


# ── WARMUP AVEC LING ──────────────────────────────────────────────────────────
def _warmup_api():
    """Warmup initial : réveille Ling et Echo au démarrage du serveur"""
    try:
        print("[BOOST] Reveil des routes Echo...")
        _requests_lib.get("http://127.0.0.1:5000/ping", timeout=2)
        
        # Warmup Ling pour la classification d'intentions
        print("[BOOST] Warmup de Ling (classification d'intentions)...")
        try:
            if client_openrouter is not None:
                warmup_ctx = {
                    "system_prompt": "Tu es un classificateur d'intentions. Réponds en JSON minimaliste.",
                    "output_tokens": 100,
                    "messages_openai": [
                        {"role": "system", "content": "Tu es un classificateur d'intentions."},
                        {"role": "user", "content": "Classifie ceci : 'comment faire un gâteau'"},
                    ],
                }
                _requests_lib.post(
                    "http://127.0.0.1:5000/horizon-warmup",
                    json={"partial": "comment"},
                    timeout=5,
                )
                print("[BOOST] Ling warmup réussi")
        except Exception as e:
            print(f"[BOOST] Ling warmup optionnel échoué: {e}")
    
    except Exception:
        pass

threading.Thread(target=_warmup_api, daemon=True).start()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)